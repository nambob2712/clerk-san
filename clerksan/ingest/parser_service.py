"""Executable, fail-closed parser sidecar.

The service accepts one request per Unix-domain connection. Parser operations
must carry exactly one read-only regular-file descriptor via ``SCM_RIGHTS``;
request paths, application credentials, and dynamic imports are never accepted.
Each untrusted parse runs in a short-lived child with hard rlimits and is killed
and reaped when its wall or output budget is exceeded.
"""

from __future__ import annotations

import argparse
import array
import datetime as dt
import fcntl
import gzip
import hashlib
import io
import json
import math
import os
import re
import resource
import select
import signal
import socket
import stat
import sys
import tarfile
import tempfile
import time
import zipfile
from collections.abc import Callable, Mapping
from dataclasses import dataclass
from pathlib import Path
from types import MappingProxyType
from typing import Any

from clerksan.ingest.filetype import MIME_BY_FILE_TYPE, FileType
from clerksan.ingest.limits import IngestLimits, ResourceLimitExceeded
from clerksan.ingest.normalized import NormalizedDocument
from clerksan.ingest.parser_artifacts import (
    MAX_ARTIFACT_FDS,
    AdapterRunResult,
    GeneratedArtifact,
    ParserArtifact,
    ParserArtifactError,
    create_sealed_artifact_fd,
    descriptor_for_generated,
    validate_received_artifacts,
    validate_result_artifact_set,
)
from clerksan.ingest.parser_runner import (
    MAX_NONCE_LENGTH,
    MAX_PROTOCOL_REQUEST_BYTES,
    MAX_PROTOCOL_RESPONSE_BYTES,
    PROTOCOL_SCHEMA,
    PROTOCOL_VERSION,
    AdapterContext,
    ReadOnlySource,
    SandboxProtocolError,
    SidecarSandboxBackend,
)

_FD_BYTES = array.array("i").itemsize
_MAX_RECEIVED_FDS = 16
_STARTUP_FIXTURE_DATETIME = dt.datetime(2000, 1, 1)
_STARTUP_FIXTURE_ZIP_TIMESTAMP = (2000, 1, 1, 0, 0, 0)
_SAFE_CHILD_ENVIRONMENT = {
    "HOME": "/nonexistent",
    "LANG": "C.UTF-8",
    "PATH": "/usr/local/bin:/usr/bin:/bin",
    "PYTHONDONTWRITEBYTECODE": "1",
    "PYTHONUNBUFFERED": "1",
}
_FORBIDDEN_REQUEST_KEYS = frozenset(
    {
        "api_key",
        "database_password",
        "database_url",
        "host_path",
        "model_path",
        "ollama_url",
        "password",
        "path",
        "secret",
        "storage_dir",
        "store_path",
        "token",
    }
)
_FORBIDDEN_ENVIRONMENT_FRAGMENTS = (
    "API_KEY",
    "DATABASE",
    "OLLAMA",
    "PASSWORD",
    "SECRET",
    "STORAGE",
    "STORE_PATH",
    "TOKEN",
)
_COMMON_REQUEST_KEYS = frozenset({"schema", "version", "nonce", "operation"})
_SOURCE_KEYS = frozenset({"filename", "mime_type", "source_id", "source_version"})
_RUN_KEYS = _COMMON_REQUEST_KEYS | frozenset(
    {
        "source_sha256",
        "source",
        "adapter_key",
        "adapter_version",
        "policy_version",
        "registry_digest",
        "metadata",
        "limits",
    }
)
_PREFLIGHT_KEYS = _COMMON_REQUEST_KEYS | frozenset(
    {
        "source_sha256",
        "source",
        "policy_version",
        "registry_digest",
        "detected",
        "limits",
    }
)
_PROBE_KEYS = _COMMON_REQUEST_KEYS


@dataclass(frozen=True, slots=True)
class ParserRequest:
    """One validated, descriptor-bound sidecar request."""

    operation: str
    nonce: str
    source_sha256: str | None
    payload: Mapping[str, Any]

    def __post_init__(self) -> None:
        object.__setattr__(self, "payload", MappingProxyType(dict(self.payload)))


@dataclass(frozen=True, slots=True)
class ParserServiceConfig:
    """Secret-free hard limits read only by the parser service."""

    socket_path: Path = Path("/run/clerksan-parser/parser.sock")
    socket_mode: int = 0o660
    wall_timeout_seconds: float = 15.0
    cpu_seconds: int = 10
    memory_bytes: int = 512 * 1024 * 1024
    file_bytes: int = 64 * 1024 * 1024
    open_files: int = 256
    processes: int = 4
    max_child_output_bytes: int = MAX_PROTOCOL_RESPONSE_BYTES - 1

    def __post_init__(self) -> None:
        path = Path(self.socket_path)
        if not path.is_absolute() or path.name in {"", ".", ".."}:
            raise ValueError("parser socket path must be absolute")
        if path.suffix != ".sock":
            raise ValueError("parser socket path must end in .sock")
        object.__setattr__(self, "socket_path", path)
        if self.socket_mode < 0o600 or self.socket_mode > 0o770 or self.socket_mode & 0o007:
            raise ValueError("parser socket mode must deny world access")
        if not math.isfinite(self.wall_timeout_seconds) or self.wall_timeout_seconds <= 0:
            raise ValueError("parser wall timeout must be greater than zero")
        for name in (
            "cpu_seconds",
            "memory_bytes",
            "file_bytes",
            "open_files",
            "processes",
            "max_child_output_bytes",
        ):
            value = getattr(self, name)
            if not isinstance(value, int) or isinstance(value, bool) or value < 1:
                raise ValueError(f"{name} must be a positive integer")
        if self.max_child_output_bytes >= MAX_PROTOCOL_RESPONSE_BYTES:
            raise ValueError("child output limit must be below the protocol response limit")

    @classmethod
    def from_environment(cls, environment: Mapping[str, str] | None = None) -> ParserServiceConfig:
        """Read only ``CLERKSAN_PARSER_*`` values, never application settings."""

        values = environment if environment is not None else os.environ
        defaults = cls()
        return cls(
            socket_path=Path(values.get("CLERKSAN_PARSER_SOCKET_PATH", str(defaults.socket_path))),
            socket_mode=int(
                values.get("CLERKSAN_PARSER_SOCKET_MODE", oct(defaults.socket_mode)), 8
            ),
            wall_timeout_seconds=float(
                values.get(
                    "CLERKSAN_PARSER_WALL_TIMEOUT_SECONDS",
                    str(defaults.wall_timeout_seconds),
                )
            ),
            cpu_seconds=int(values.get("CLERKSAN_PARSER_CPU_SECONDS", str(defaults.cpu_seconds))),
            memory_bytes=int(
                values.get("CLERKSAN_PARSER_MEMORY_BYTES", str(defaults.memory_bytes))
            ),
            file_bytes=int(values.get("CLERKSAN_PARSER_FILE_BYTES", str(defaults.file_bytes))),
            open_files=int(values.get("CLERKSAN_PARSER_OPEN_FILES", str(defaults.open_files))),
            processes=int(values.get("CLERKSAN_PARSER_PROCESSES", str(defaults.processes))),
            max_child_output_bytes=int(
                values.get(
                    "CLERKSAN_PARSER_MAX_OUTPUT_BYTES",
                    str(defaults.max_child_output_bytes),
                )
            ),
        )

    def evidence_payload(self) -> dict[str, int | float | str]:
        """Return bounded, non-secret settings that participate in probe evidence."""

        return {
            "socket_mode": oct(self.socket_mode),
            "wall_timeout_seconds": self.wall_timeout_seconds,
            "cpu_seconds": self.cpu_seconds,
            "memory_bytes": self.memory_bytes,
            "file_bytes": self.file_bytes,
            "open_files": self.open_files,
            "processes": self.processes,
            "max_child_output_bytes": self.max_child_output_bytes,
        }


RunHandler = Callable[
    [ReadOnlySource, AdapterContext, IngestLimits],
    NormalizedDocument | AdapterRunResult | Mapping[str, Any],
]
PreflightHandler = Callable[[ReadOnlySource, Mapping[str, Any], IngestLimits], Mapping[str, Any]]


@dataclass(frozen=True, slots=True)
class _AdapterEntry:
    adapter_key: str
    format_key: str
    adapter_version: str
    run: RunHandler
    preflight: PreflightHandler
    startup_fixture: Callable[[], bytes] | None


@dataclass(frozen=True, slots=True)
class StartupSelfTestResult:
    format_key: str
    passed: bool
    evidence_digest: str | None = None


class ParserAdapterRegistry:
    """Closed parser dispatch table; no request-driven imports or callables."""

    def __init__(self) -> None:
        self._by_adapter: dict[str, _AdapterEntry] = {}
        self._by_format: dict[str, _AdapterEntry] = {}

    def register(
        self,
        *,
        adapter_key: str,
        format_key: str,
        run: RunHandler,
        preflight: PreflightHandler,
        adapter_version: str = "1",
        startup_fixture: Callable[[], bytes] | None = None,
    ) -> None:
        for name, value in (
            ("adapter_key", adapter_key),
            ("format_key", format_key),
            ("adapter_version", adapter_version),
        ):
            _require_stable_token(name, value)
        if adapter_key != adapter_key.casefold() or format_key != format_key.casefold():
            raise ValueError("parser adapter and format keys must be lowercase")
        if adapter_key in self._by_adapter:
            raise ValueError(f"duplicate parser adapter key: {adapter_key}")
        if format_key in self._by_format:
            raise ValueError(f"duplicate parser format key: {format_key}")
        entry = _AdapterEntry(
            adapter_key,
            format_key,
            adapter_version,
            run,
            preflight,
            startup_fixture,
        )
        self._by_adapter[adapter_key] = entry
        self._by_format[format_key] = entry

    @property
    def capabilities(self) -> tuple[str, ...]:
        return tuple(sorted(self._by_format))

    @property
    def registry_digest(self) -> str:
        return self.registry_digest_for(self.capabilities)

    def registry_digest_for(self, formats: tuple[str, ...] | frozenset[str]) -> str:
        selected = frozenset(formats)
        payload = {
            "schema": "clerksan.parser-registry",
            "version": 1,
            "protocol_version": PROTOCOL_VERSION,
            "adapters": [
                {
                    "adapter_key": entry.adapter_key,
                    "format_key": entry.format_key,
                    "adapter_version": entry.adapter_version,
                }
                for entry in sorted(self._by_adapter.values(), key=lambda item: item.adapter_key)
                if entry.format_key in selected
            ],
        }
        return hashlib.sha256(_canonical_json(payload)).hexdigest()

    def format_for_adapter(self, adapter_key: str) -> str:
        try:
            return self._by_adapter[adapter_key].format_key
        except KeyError as error:
            raise LookupError("parser adapter is not registered") from error

    def startup_self_test(
        self,
        limits: IngestLimits | None = None,
        *,
        config: ParserServiceConfig | None = None,
    ) -> tuple[StartupSelfTestResult, ...]:
        """Parse one built-in bounded fixture for every registered format."""

        active_limits = limits or IngestLimits()
        active_config = config or ParserServiceConfig()
        return tuple(
            _run_isolated_startup_self_test(entry, active_limits, active_config)
            for entry in sorted(self._by_adapter.values(), key=lambda item: item.format_key)
        )

    def run(
        self,
        adapter_key: str,
        adapter_version: str,
        source: ReadOnlySource,
        context: AdapterContext,
        limits: IngestLimits,
    ) -> NormalizedDocument | AdapterRunResult | Mapping[str, Any]:
        try:
            entry = self._by_adapter[adapter_key]
        except KeyError as error:
            raise LookupError("parser adapter is not registered") from error
        if adapter_version != entry.adapter_version:
            raise LookupError("parser adapter version is not registered")
        return entry.run(source, context, limits)

    def preflight(
        self,
        source: ReadOnlySource,
        detected: Mapping[str, Any],
        limits: IngestLimits,
    ) -> Mapping[str, Any]:
        format_key = detected.get("format")
        if not isinstance(format_key, str):
            raise ValueError("detected format is required")
        try:
            entry = self._by_format[format_key]
        except KeyError as error:
            raise LookupError("parser preflight is not registered") from error
        return entry.preflight(source, detected, limits)


@dataclass(frozen=True, slots=True)
class SandboxRuntimeEvidence:
    """Direct observations made inside the parser container."""

    network_isolated: bool
    secrets_absent: bool
    root_read_only: bool
    capabilities_dropped: bool
    no_new_privileges: bool
    tmpfs_hardened: bool
    cgroup_bounded: bool
    child_timeout_reaped: bool

    @property
    def verified(self) -> bool:
        return all(self.as_dict().values())

    def as_dict(self) -> dict[str, bool]:
        return {
            "network_isolated": self.network_isolated,
            "secrets_absent": self.secrets_absent,
            "root_read_only": self.root_read_only,
            "capabilities_dropped": self.capabilities_dropped,
            "no_new_privileges": self.no_new_privileges,
            "tmpfs_hardened": self.tmpfs_hardened,
            "cgroup_bounded": self.cgroup_bounded,
            "child_timeout_reaped": self.child_timeout_reaped,
        }


@dataclass(frozen=True, slots=True)
class ChildOutcome:
    payload: Mapping[str, Any] | None
    reason: str | None
    reaped: bool


@dataclass(frozen=True, slots=True)
class ChildArtifactOutcome:
    payload: Mapping[str, Any] | None
    reason: str | None
    reaped: bool
    artifact_fds: tuple[int, ...] = ()


@dataclass(frozen=True, slots=True)
class _DispatchResult:
    payload: Mapping[str, Any]
    artifacts: tuple[GeneratedArtifact, ...] = ()


def validate_request(
    payload: Mapping[str, Any], *, has_source_fd: bool | None = None, fd_count: int | None = None
) -> ParserRequest:
    """Validate a decoded request before any parser is selected."""

    if not isinstance(payload, Mapping):
        raise SandboxProtocolError("parser request must be an object")
    if payload.get("schema") != PROTOCOL_SCHEMA or payload.get("version") != PROTOCOL_VERSION:
        raise SandboxProtocolError("unsupported parser request schema")
    nonce = payload.get("nonce")
    if not isinstance(nonce, str) or len(nonce) < 16 or len(nonce) > MAX_NONCE_LENGTH:
        raise SandboxProtocolError("invalid parser request nonce")
    operation = payload.get("operation")
    if operation not in {"probe", "preflight", "run"}:
        raise SandboxProtocolError("unsupported parser operation")

    if fd_count is None:
        if has_source_fd is None:
            raise SandboxProtocolError("parser descriptor count is required")
        fd_count = 1 if has_source_fd else 0
    if not isinstance(fd_count, int) or isinstance(fd_count, bool) or fd_count < 0:
        raise SandboxProtocolError("invalid parser descriptor count")
    expected_fd_count = 0 if operation == "probe" else 1
    if fd_count != expected_fd_count:
        expected = "no" if expected_fd_count == 0 else "exactly one"
        raise SandboxProtocolError(f"parser operation requires {expected} source descriptor")

    _reject_forbidden_request_data(payload)
    allowed_keys = {
        "probe": _PROBE_KEYS,
        "preflight": _PREFLIGHT_KEYS,
        "run": _RUN_KEYS,
    }[operation]
    unknown = set(payload) - allowed_keys
    if unknown:
        raise SandboxProtocolError("parser request contains unsupported fields")

    source_sha256 = payload.get("source_sha256")
    if operation == "probe":
        source_sha256 = None
    else:
        if not _is_sha256(source_sha256):
            raise SandboxProtocolError("parser operation requires source digest")
        _validate_source_metadata(payload.get("source"))
        _parse_limits(payload.get("limits"))
        _require_stable_token("policy_version", payload.get("policy_version"))
        if not _is_sha256(payload.get("registry_digest")):
            raise SandboxProtocolError("parser operation requires registry digest")

    if operation == "preflight":
        detected = payload.get("detected")
        if not isinstance(detected, Mapping) or not isinstance(detected.get("format"), str):
            raise SandboxProtocolError("preflight requires detected format evidence")
        _require_bounded_json("detected evidence", detected, 16 * 1024)
    elif operation == "run":
        _require_stable_token("adapter_key", payload.get("adapter_key"))
        _require_stable_token("adapter_version", payload.get("adapter_version"))
        metadata = payload.get("metadata")
        if not isinstance(metadata, Mapping) or len(metadata) > 64:
            raise SandboxProtocolError("adapter metadata must be a bounded object")
        if not all(
            isinstance(key, str)
            and len(key) <= 128
            and isinstance(value, str)
            and len(value) <= 4096
            for key, value in metadata.items()
        ):
            raise SandboxProtocolError("adapter metadata must contain bounded strings")
        _require_bounded_json("adapter metadata", metadata, 32 * 1024)

    return ParserRequest(operation, nonce, source_sha256, dict(payload))


def response(*, nonce: str, source_sha256: str | None, ok: bool, **payload: Any) -> dict[str, Any]:
    """Build a canonical response whose binding fields cannot be overridden."""

    if not isinstance(nonce, str) or not 16 <= len(nonce) <= MAX_NONCE_LENGTH:
        raise ValueError("nonce must be a bounded string")
    if source_sha256 is not None and not _is_sha256(source_sha256):
        raise ValueError("source_sha256 must be a lowercase SHA-256 digest")
    reserved = {"schema", "version", "nonce", "source_sha256", "ok"}
    if reserved.intersection(payload):
        raise ValueError("response payload cannot override protocol bindings")
    result = dict(payload)
    result.update(
        {
            "schema": PROTOCOL_SCHEMA,
            "version": PROTOCOL_VERSION,
            "nonce": nonce,
            "source_sha256": source_sha256,
            "ok": bool(ok),
        }
    )
    return result


def build_default_registry() -> ParserAdapterRegistry:
    """Build the audited sidecar registry without discovering plugins dynamically."""

    from clerksan.ingest.adapters.archive import ArchiveAdapter
    from clerksan.ingest.adapters.docx import DocxAdapter
    from clerksan.ingest.adapters.email import EmailAdapter
    from clerksan.ingest.adapters.html import HtmlAdapter
    from clerksan.ingest.adapters.image import ImageAdapter
    from clerksan.ingest.adapters.markdown import MarkdownAdapter
    from clerksan.ingest.adapters.odf import OdfAdapter
    from clerksan.ingest.adapters.pdf import PdfAdapter
    from clerksan.ingest.adapters.pptx import PptxAdapter
    from clerksan.ingest.adapters.rtf import RtfAdapter
    from clerksan.ingest.adapters.structured import StructuredAdapter
    from clerksan.ingest.adapters.text import TextAdapter
    from clerksan.ingest.adapters.xlsx import XlsxAdapter

    registry = ParserAdapterRegistry()
    factories: dict[str, Callable[[IngestLimits], Any]] = {
        "md": lambda limits: MarkdownAdapter(limits=limits),
        "txt": lambda limits: TextAdapter(limits=limits),
        "rst": lambda limits: TextAdapter(limits=limits),
        "log": lambda limits: TextAdapter(limits=limits),
        "json": lambda limits: StructuredAdapter(limits=limits),
        "jsonl": lambda limits: StructuredAdapter(limits=limits),
        "yaml": lambda limits: StructuredAdapter(limits=limits),
        "xml": lambda limits: StructuredAdapter(limits=limits),
        "svg": lambda limits: StructuredAdapter(limits=limits),
        "html": lambda limits: HtmlAdapter(limits=limits),
        "docx": lambda limits: DocxAdapter(limits=limits),
        "xlsx": lambda limits: XlsxAdapter(limits=limits),
        "pptx": lambda limits: PptxAdapter(limits=limits),
        "odt": lambda limits: OdfAdapter(limits=limits),
        "odp": lambda limits: OdfAdapter(limits=limits),
        "ods": lambda limits: OdfAdapter(limits=limits),
        "rtf": lambda limits: RtfAdapter(limits=limits),
        "eml": lambda limits: EmailAdapter(limits=limits),
        "zip": lambda limits: ArchiveAdapter(limits=limits),
        "tar": lambda limits: ArchiveAdapter(limits=limits),
        "tgz": lambda limits: ArchiveAdapter(limits=limits),
        "gz": lambda limits: ArchiveAdapter(limits=limits),
        "jpeg": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "png": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "webp": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "bmp": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "gif": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "tiff": lambda limits: ImageAdapter(_NoOcr(), limits=limits),
        "pdf": lambda limits: PdfAdapter(
            _NoOcr(),
            _ParserPdfPolicy(),
            limits=limits,
        ),
    }
    for format_key, factory in factories.items():
        handler = _adapter_handler(format_key, factory)
        registry.register(
            adapter_key=format_key,
            format_key=format_key,
            run=handler,
            preflight=_normalizing_preflight(format_key, handler),
            startup_fixture=lambda key=format_key: _startup_fixture(key),
        )
    for format_key in ("csv", "tsv"):
        handler = _delimited_handler(format_key)
        registry.register(
            adapter_key=f"delimited.{format_key}",
            format_key=format_key,
            run=handler,
            preflight=_normalizing_preflight(format_key, handler),
            startup_fixture=lambda key=format_key: _startup_fixture(key),
        )
    return registry


class ParserSidecarServer:
    """Single-purpose Unix socket server for parser requests."""

    def __init__(
        self,
        config: ParserServiceConfig | None = None,
        registry: ParserAdapterRegistry | None = None,
        *,
        runtime_evidence: SandboxRuntimeEvidence | None = None,
    ) -> None:
        self.config = config or ParserServiceConfig.from_environment()
        self.registry = registry or build_default_registry()
        self.startup_self_tests = self.registry.startup_self_test(config=self.config)
        self.enabled_formats = tuple(
            result.format_key for result in self.startup_self_tests if result.passed
        )
        self.registry_digest = self.registry.registry_digest_for(frozenset(self.enabled_formats))
        self.runtime_evidence = runtime_evidence or collect_runtime_evidence(self.config)
        self._listener: socket.socket | None = None
        self._stopping = False

    def serve_forever(self) -> None:
        listener = self._bind_listener()
        self._install_signal_handlers()
        try:
            while not self._stopping:
                try:
                    connection, _ = listener.accept()
                except TimeoutError:
                    continue
                with connection:
                    try:
                        self.handle_connection(connection)
                    except (OSError, SandboxProtocolError):
                        continue
        finally:
            self.close()

    def serve_once(self) -> None:
        listener = self._bind_listener()
        try:
            connection, _ = listener.accept()
            with connection:
                self.handle_connection(connection)
        finally:
            self.close()

    def handle_connection(self, connection: socket.socket) -> None:
        source_fd: int | None = None
        artifact_fds: tuple[int, ...] = ()
        request: ParserRequest | None = None
        try:
            request, source_fd = receive_request(connection)
            if request.operation == "probe":
                result = self._probe_response(request)
            else:
                assert source_fd is not None
                if not self._sandbox_ready():
                    result = response(
                        nonce=request.nonce,
                        source_sha256=request.source_sha256,
                        ok=False,
                        reason="sandbox_unavailable",
                    )
                    _send_response(connection, result)
                    return
                limits = _parse_limits(request.payload["limits"])
                _validate_source_descriptor(source_fd, limits)
                outcome = _execute_parser_child(
                    lambda: self._dispatch(request, source_fd),
                    self.config,
                    source_fd=source_fd,
                    limits=limits,
                )
                artifact_fds = outcome.artifact_fds
                if outcome.payload is None:
                    result = response(
                        nonce=request.nonce,
                        source_sha256=request.source_sha256,
                        ok=False,
                        reason=outcome.reason or "parser_child_failed",
                    )
                else:
                    result = dict(outcome.payload)
                    if request.operation == "run":
                        try:
                            normalized = NormalizedDocument.model_validate(result.get("normalized"))
                            artifacts = validate_received_artifacts(
                                result.get("artifacts", []),
                                artifact_fds,
                                source_sha256=str(request.source_sha256),
                                limits=limits,
                            )
                            validate_result_artifact_set(
                                str(request.payload["adapter_key"]),
                                normalized,
                                artifacts,
                                source_sha256=str(request.source_sha256),
                            )
                            normalized_size = len(
                                _canonical_json(normalized.model_dump(mode="json"))
                            )
                            artifact_size = sum(
                                artifact.descriptor.byte_size for artifact in artifacts
                            )
                            if normalized_size + artifact_size > limits.max_normalized_output_bytes:
                                raise ParserArtifactError(
                                    "aggregate parser output exceeds configured limit"
                                )
                        except (OSError, ParserArtifactError, TypeError, ValueError):
                            _close_descriptors(artifact_fds)
                            artifact_fds = ()
                            result = response(
                                nonce=request.nonce,
                                source_sha256=request.source_sha256,
                                ok=False,
                                reason="artifact_validation_failed",
                            )
                    elif artifact_fds:
                        _close_descriptors(artifact_fds)
                        artifact_fds = ()
                        result = response(
                            nonce=request.nonce,
                            source_sha256=request.source_sha256,
                            ok=False,
                            reason="artifact_validation_failed",
                        )
            _send_response(connection, result, artifact_fds)
        except SandboxProtocolError:
            nonce = request.nonce if request is not None else "0" * 16
            source_sha256 = request.source_sha256 if request is not None else None
            _send_response(
                connection,
                response(
                    nonce=nonce,
                    source_sha256=source_sha256,
                    ok=False,
                    reason="invalid_request",
                ),
            )
        finally:
            _close_descriptors(artifact_fds)
            if source_fd is not None:
                os.close(source_fd)

    def close(self) -> None:
        listener, self._listener = self._listener, None
        if listener is not None:
            listener.close()
        _unlink_owned_socket(self.config.socket_path)

    def _probe_response(self, request: ParserRequest) -> dict[str, Any]:
        socket_permissioned = _socket_is_permissioned(
            self.config.socket_path, self.config.socket_mode
        )
        verified = self._sandbox_ready()
        evidence = {
            **self.runtime_evidence.as_dict(),
            "socket_permissioned": socket_permissioned,
            "format_self_tests_passed": all(result.passed for result in self.startup_self_tests),
        }
        self_test_evidence = {
            result.format_key: {
                "passed": result.passed,
                "evidence_digest": result.evidence_digest,
            }
            for result in self.startup_self_tests
        }
        digest_payload = {
            "schema": "clerksan.parser-sandbox-evidence",
            "version": 1,
            "runtime": evidence,
            "limits": self.config.evidence_payload(),
            "registry_digest": self.registry_digest,
            "capabilities": list(self.enabled_formats),
            "format_self_tests": self_test_evidence,
        }
        return response(
            nonce=request.nonce,
            source_sha256=None,
            ok=True,
            verified=verified,
            backend="parser-sidecar" if verified else None,
            evidence_digest=hashlib.sha256(_canonical_json(digest_payload)).hexdigest()
            if verified
            else None,
            registry_digest=self.registry_digest,
            capabilities=list(self.enabled_formats) if verified else [],
            evidence=evidence,
            format_self_tests=self_test_evidence,
        )

    def _sandbox_ready(self) -> bool:
        return self.runtime_evidence.verified and _socket_is_permissioned(
            self.config.socket_path, self.config.socket_mode
        )

    def _dispatch(self, request: ParserRequest, source_fd: int) -> _DispatchResult:
        try:
            if request.payload["registry_digest"] != self.registry_digest:
                raise SandboxProtocolError("registry digest mismatch")
            limits = _parse_limits(request.payload["limits"])
            source = _source_from_request(source_fd, request)
            _validate_source_descriptor(source_fd, limits)
            source.verify_digest(max_bytes=limits.max_upload_bytes)
            if request.operation == "preflight":
                if request.payload["detected"]["format"] not in self.enabled_formats:
                    raise LookupError("parser preflight self-test did not pass")
                evidence = self.registry.preflight(source, request.payload["detected"], limits)
                return _DispatchResult(
                    response(
                        nonce=request.nonce,
                        source_sha256=request.source_sha256,
                        ok=True,
                        evidence={"schema_version": 1, **dict(evidence)},
                    )
                )

            adapter_key = str(request.payload["adapter_key"])
            if self.registry.format_for_adapter(adapter_key) not in self.enabled_formats:
                raise LookupError("parser adapter self-test did not pass")
            context = AdapterContext(
                adapter_key=adapter_key,
                adapter_version=str(request.payload["adapter_version"]),
                policy_version=str(request.payload["policy_version"]),
                registry_digest=str(request.payload["registry_digest"]),
                metadata=request.payload["metadata"],
            )
            adapter_result = self.registry.run(
                adapter_key,
                context.adapter_version,
                source,
                context,
                limits,
            )
            artifacts: tuple[GeneratedArtifact, ...]
            if isinstance(adapter_result, AdapterRunResult):
                normalized = adapter_result.normalized
                artifacts = adapter_result.artifacts
            else:
                normalized = adapter_result
                artifacts = ()
            if isinstance(normalized, NormalizedDocument):
                normalized_payload = normalized.model_dump(mode="json")
            elif isinstance(normalized, Mapping):
                normalized_payload = dict(normalized)
                NormalizedDocument.model_validate(normalized_payload)
            else:
                raise TypeError("parser returned an unsupported result")
            _require_bounded_json(
                "normalized output", normalized_payload, limits.max_normalized_output_bytes
            )
            normalized_size = len(_canonical_json(normalized_payload))
            artifact_size = sum(len(artifact.data) for artifact in artifacts)
            if normalized_size + artifact_size > limits.max_normalized_output_bytes:
                raise ResourceLimitExceeded(
                    "max_normalized_output_bytes",
                    limits.max_normalized_output_bytes,
                    normalized_size + artifact_size,
                )
            return _DispatchResult(
                response(
                    nonce=request.nonce,
                    source_sha256=request.source_sha256,
                    ok=True,
                    normalized=normalized_payload,
                ),
                artifacts,
            )
        except LookupError:
            return _DispatchResult(
                response(
                    nonce=request.nonce,
                    source_sha256=request.source_sha256,
                    ok=False,
                    reason="adapter_unavailable",
                )
            )
        except (MemoryError, ResourceLimitExceeded):
            return _DispatchResult(
                response(
                    nonce=request.nonce,
                    source_sha256=request.source_sha256,
                    ok=False,
                    reason="resource_limit_exceeded",
                )
            )
        except SandboxProtocolError:
            return _DispatchResult(
                response(
                    nonce=request.nonce,
                    source_sha256=request.source_sha256,
                    ok=False,
                    reason="protocol_binding_failed",
                )
            )
        except (OSError, TypeError, ValueError):
            return _DispatchResult(
                response(
                    nonce=request.nonce,
                    source_sha256=request.source_sha256,
                    ok=False,
                    reason="parser_failed",
                )
            )

    def _bind_listener(self) -> socket.socket:
        if self._listener is not None:
            return self._listener
        socket_path = self.config.socket_path
        parent = socket_path.parent
        parent.mkdir(mode=0o750, parents=True, exist_ok=True)
        parent_stat = os.lstat(parent)
        if stat.S_ISLNK(parent_stat.st_mode) or parent_stat.st_mode & 0o002:
            raise RuntimeError("parser socket directory is not trusted")
        _unlink_owned_socket(socket_path)
        listener = socket.socket(socket.AF_UNIX, socket.SOCK_STREAM)
        try:
            listener.bind(str(socket_path))
            os.chmod(socket_path, self.config.socket_mode)
            listener.listen(16)
            listener.settimeout(0.5)
        except BaseException:
            listener.close()
            _unlink_owned_socket(socket_path)
            raise
        self._listener = listener
        return listener

    def _install_signal_handlers(self) -> None:
        def stop(_signum: int, _frame: object) -> None:
            self._stopping = True

        signal.signal(signal.SIGTERM, stop)
        signal.signal(signal.SIGINT, stop)


def receive_request(connection: socket.socket) -> tuple[ParserRequest, int | None]:
    """Receive one bounded JSON line and all attached descriptors."""

    data = bytearray()
    descriptors: list[int] = []
    try:
        while b"\n" not in data:
            chunk, ancillary, flags, _ = connection.recvmsg(
                min(65536, MAX_PROTOCOL_REQUEST_BYTES - len(data) + 1),
                socket.CMSG_SPACE(_FD_BYTES * _MAX_RECEIVED_FDS),
            )
            if not chunk:
                break
            data.extend(chunk)
            descriptors.extend(_descriptors_from_ancillary(ancillary))
            if flags & socket.MSG_CTRUNC:
                raise SandboxProtocolError("parser descriptor control data was truncated")
            if len(data) > MAX_PROTOCOL_REQUEST_BYTES:
                raise SandboxProtocolError("parser request exceeds protocol limit")
            if len(descriptors) > _MAX_RECEIVED_FDS:
                raise SandboxProtocolError("too many parser descriptors")
        if b"\n" not in data:
            raise SandboxProtocolError("parser request is not newline terminated")
        encoded, trailing = bytes(data).split(b"\n", 1)
        if trailing:
            raise SandboxProtocolError("parser request contains trailing data")
        try:
            payload = json.loads(encoded)
        except (UnicodeDecodeError, ValueError) as error:
            raise SandboxProtocolError("parser request is not valid JSON") from error
        request = validate_request(payload, fd_count=len(descriptors))
        for descriptor in descriptors:
            fcntl.fcntl(descriptor, fcntl.F_SETFD, fcntl.FD_CLOEXEC)
        return request, descriptors[0] if descriptors else None
    except BaseException:
        for descriptor in descriptors:
            os.close(descriptor)
        raise


def collect_runtime_evidence(config: ParserServiceConfig) -> SandboxRuntimeEvidence:
    """Inspect the live container and mechanically prove child timeout/reaping."""

    timeout_probe = _execute_child(
        lambda: _sleep_then_payload(1.0),
        config,
        wall_timeout_seconds=min(0.05, config.wall_timeout_seconds),
    )
    return SandboxRuntimeEvidence(
        network_isolated=_network_namespace_isolated(),
        secrets_absent=_application_secrets_absent(),
        root_read_only=_mount_is_read_only("/"),
        capabilities_dropped=_status_value("CapEff") == "0000000000000000",
        no_new_privileges=_status_value("NoNewPrivs") == "1",
        tmpfs_hardened=_tmpfs_is_hardened("/tmp"),
        cgroup_bounded=_cgroup_is_bounded(config),
        child_timeout_reaped=(timeout_probe.reason == "parser_timeout" and timeout_probe.reaped),
    )


def _execute_child(
    task: Callable[[], Mapping[str, Any]],
    config: ParserServiceConfig,
    *,
    source_fd: int | None = None,
    wall_timeout_seconds: float | None = None,
) -> ChildOutcome:
    """Run one task in a limited process and always reap it."""

    read_fd, write_fd = os.pipe()
    pid = os.fork()
    if pid == 0:
        try:
            os.close(read_fd)
            os.setsid()
            _apply_child_limits(config)
            os.environ.clear()
            os.environ.update(_SAFE_CHILD_ENVIRONMENT)
            keep = {0, 1, 2, write_fd}
            if source_fd is not None:
                keep.add(source_fd)
            _close_unneeded_descriptors(keep)
            try:
                payload = dict(task())
                encoded = _canonical_json(payload) + b"\n"
            except BaseException:
                encoded = b'{"internal_child_error":true}\n'
            if len(encoded) > config.max_child_output_bytes:
                encoded = b'{"child_output_exceeded":true}\n'
            _write_all(write_fd, encoded)
            os.close(write_fd)
            os._exit(0)
        except BaseException:
            os._exit(127)

    os.close(write_fd)
    os.set_blocking(read_fd, False)
    timeout = wall_timeout_seconds or config.wall_timeout_seconds
    deadline = time.monotonic() + timeout
    output = bytearray()
    status: int | None = None
    reaped = False
    reason: str | None = None
    try:
        while True:
            remaining = deadline - time.monotonic()
            if remaining <= 0:
                reason = "parser_timeout"
                reaped = _kill_and_reap(pid)
                break
            readable, _, _ = select.select([read_fd], [], [], min(remaining, 0.05))
            if readable:
                while True:
                    try:
                        chunk = os.read(read_fd, 65536)
                    except BlockingIOError:
                        break
                    if not chunk:
                        break
                    output.extend(chunk)
                    if len(output) > config.max_child_output_bytes:
                        reason = "parser_output_limit"
                        reaped = _kill_and_reap(pid)
                        break
                if reason is not None:
                    break
            waited_pid, child_status = os.waitpid(pid, os.WNOHANG)
            if waited_pid == pid:
                status = child_status
                reaped = True
                while True:
                    try:
                        chunk = os.read(read_fd, 65536)
                    except BlockingIOError:
                        break
                    if not chunk:
                        break
                    output.extend(chunk)
                break
        if reason is not None:
            return ChildOutcome(None, reason, reaped)
        if status is not None and os.WIFSIGNALED(status):
            resource_signals = {
                signal.SIGKILL,
                getattr(signal, "SIGXCPU", signal.SIGKILL),
                getattr(signal, "SIGXFSZ", signal.SIGKILL),
            }
            signal_reason = (
                "parser_resource_limit"
                if os.WTERMSIG(status) in resource_signals
                else "parser_child_failed"
            )
            return ChildOutcome(None, signal_reason, reaped)
        if status is None or not os.WIFEXITED(status) or os.WEXITSTATUS(status) != 0:
            return ChildOutcome(None, "parser_child_failed", reaped)
        if len(output) > config.max_child_output_bytes:
            return ChildOutcome(None, "parser_output_limit", reaped)
        if not output.endswith(b"\n") or output.count(b"\n") != 1:
            return ChildOutcome(None, "parser_child_failed", reaped)
        try:
            payload = json.loads(output[:-1])
        except (UnicodeDecodeError, ValueError):
            return ChildOutcome(None, "parser_child_failed", reaped)
        if not isinstance(payload, Mapping) or payload.get("internal_child_error") is True:
            return ChildOutcome(None, "parser_child_failed", reaped)
        if payload.get("child_output_exceeded") is True:
            return ChildOutcome(None, "parser_output_limit", reaped)
        return ChildOutcome(dict(payload), None, reaped)
    finally:
        os.close(read_fd)
        if not reaped:
            _kill_and_reap(pid)


def _execute_parser_child(
    task: Callable[[], _DispatchResult],
    config: ParserServiceConfig,
    *,
    source_fd: int,
    limits: IngestLimits,
) -> ChildArtifactOutcome:
    """Run a parser child and receive metadata plus sealed output descriptors."""

    parent_socket, child_socket = socket.socketpair(socket.AF_UNIX, socket.SOCK_STREAM)
    pid = os.fork()
    if pid == 0:
        output_fds: list[int] = []
        try:
            parent_socket.close()
            os.setsid()
            _apply_child_limits(config)
            os.environ.clear()
            os.environ.update(_SAFE_CHILD_ENVIRONMENT)
            _close_unneeded_descriptors({0, 1, 2, source_fd, child_socket.fileno()})
            dispatch = task()
            artifacts = tuple(dispatch.artifacts)
            if len(artifacts) > min(MAX_ARTIFACT_FDS, limits.max_pdf_pages + 1):
                raise ParserArtifactError("artifact descriptor count exceeds configured limit")
            if sum(len(artifact.data) for artifact in artifacts) > (
                limits.max_normalized_output_bytes
            ):
                raise ParserArtifactError("artifact output exceeds configured limit")
            metadata: list[dict[str, Any]] = []
            source_sha256 = dispatch.payload.get("source_sha256")
            if artifacts and not _is_sha256(source_sha256):
                raise ParserArtifactError("artifact response lacks source binding")
            for ordinal, artifact in enumerate(artifacts, start=1):
                artifact_fd, seal_supported, sealed = create_sealed_artifact_fd(artifact)
                output_fds.append(artifact_fd)
                metadata.append(
                    descriptor_for_generated(
                        artifact,
                        ordinal=ordinal,
                        source_sha256=str(source_sha256),
                        seal_supported=seal_supported,
                        sealed=sealed,
                    ).as_dict()
                )
            payload = {**dict(dispatch.payload), "artifacts": metadata}
            encoded = _canonical_json(payload) + b"\n"
            if len(encoded) > config.max_child_output_bytes:
                raise ParserArtifactError("parser child control output exceeds configured limit")
            _send_encoded_response(child_socket, encoded, output_fds)
            for descriptor in output_fds:
                os.close(descriptor)
            child_socket.close()
            os._exit(0)
        except BaseException:
            _close_descriptors(output_fds)
            try:
                child_socket.sendall(b'{"internal_child_error":true}\n')
            except OSError:
                pass
            child_socket.close()
            os._exit(127)

    child_socket.close()
    parent_socket.settimeout(config.wall_timeout_seconds)
    artifact_fds: tuple[int, ...] = ()
    reaped = False
    try:
        try:
            payload, artifact_fds = _receive_child_response(
                parent_socket,
                maximum=config.max_child_output_bytes,
            )
        except TimeoutError:
            reaped = _kill_and_reap(pid)
            return ChildArtifactOutcome(
                None,
                "parser_timeout",
                reaped,
            )
        except (OSError, SandboxProtocolError):
            reaped = _wait_or_kill_child(pid, time.monotonic() + 0.2)
            return ChildArtifactOutcome(None, "parser_child_failed", reaped)

        reaped, status = _wait_for_child_status(
            pid,
            deadline=time.monotonic() + config.wall_timeout_seconds,
        )
        if not reaped:
            _close_descriptors(artifact_fds)
            artifact_fds = ()
            reaped = _kill_and_reap(pid)
            return ChildArtifactOutcome(None, "parser_timeout", reaped)
        if status is not None and os.WIFSIGNALED(status):
            resource_signals = {
                signal.SIGKILL,
                getattr(signal, "SIGXCPU", signal.SIGKILL),
                getattr(signal, "SIGXFSZ", signal.SIGKILL),
            }
            reason = (
                "parser_resource_limit"
                if os.WTERMSIG(status) in resource_signals
                else "parser_child_failed"
            )
            _close_descriptors(artifact_fds)
            artifact_fds = ()
            return ChildArtifactOutcome(None, reason, True)
        if (
            status is None
            or not os.WIFEXITED(status)
            or os.WEXITSTATUS(status) != 0
            or payload.get("internal_child_error") is True
        ):
            _close_descriptors(artifact_fds)
            artifact_fds = ()
            return ChildArtifactOutcome(None, "parser_child_failed", True)
        return ChildArtifactOutcome(payload, None, True, artifact_fds)
    finally:
        parent_socket.close()
        if not reaped:
            _kill_and_reap(pid)


def _receive_child_response(
    connection: socket.socket,
    *,
    maximum: int,
) -> tuple[dict[str, Any], tuple[int, ...]]:
    data = bytearray()
    descriptors: list[int] = []
    try:
        while b"\n" not in data:
            chunk, ancillary, flags, _ = connection.recvmsg(
                min(65536, maximum - len(data) + 1),
                socket.CMSG_SPACE(_FD_BYTES * MAX_ARTIFACT_FDS),
            )
            descriptors.extend(_descriptors_from_ancillary(ancillary))
            if flags & socket.MSG_CTRUNC:
                raise SandboxProtocolError("artifact control data was truncated")
            if len(descriptors) > MAX_ARTIFACT_FDS:
                raise SandboxProtocolError("too many child artifact descriptors")
            if not chunk:
                break
            data.extend(chunk)
            if len(data) > maximum:
                raise SandboxProtocolError("parser child output exceeds configured limit")
        if b"\n" not in data:
            raise SandboxProtocolError("parser child response is not newline terminated")
        encoded, trailing = bytes(data).split(b"\n", 1)
        if trailing:
            raise SandboxProtocolError("parser child response contains trailing data")
        try:
            payload = json.loads(encoded)
        except (UnicodeDecodeError, ValueError) as error:
            raise SandboxProtocolError("parser child response is invalid JSON") from error
        if not isinstance(payload, dict):
            raise SandboxProtocolError("parser child response must be an object")
        for descriptor in descriptors:
            fcntl.fcntl(descriptor, fcntl.F_SETFD, fcntl.FD_CLOEXEC)
        return payload, tuple(descriptors)
    except BaseException:
        _close_descriptors(descriptors)
        raise


def _wait_for_child_status(pid: int, *, deadline: float) -> tuple[bool, int | None]:
    while time.monotonic() < deadline:
        waited_pid, status = os.waitpid(pid, os.WNOHANG)
        if waited_pid == pid:
            return True, status
        time.sleep(0.001)
    return False, None


def _wait_or_kill_child(pid: int, deadline: float) -> bool:
    reaped, _ = _wait_for_child_status(pid, deadline=deadline)
    return reaped or _kill_and_reap(pid)


def _parse_limits(value: object) -> IngestLimits:
    if not isinstance(value, Mapping):
        raise SandboxProtocolError("parser limits must be an object")
    fields = IngestLimits.__dataclass_fields__
    if set(value) != set(fields):
        raise SandboxProtocolError("parser limits must match the versioned limit schema")
    parsed: dict[str, int | float] = {}
    for name in fields:
        item = value[name]
        if name == "max_archive_expansion_ratio":
            if (
                not isinstance(item, (int, float))
                or isinstance(item, bool)
                or not math.isfinite(float(item))
                or item < 1
            ):
                raise SandboxProtocolError("invalid parser expansion ratio")
            parsed[name] = float(item)
        else:
            if not isinstance(item, int) or isinstance(item, bool) or item < 1:
                raise SandboxProtocolError("invalid parser integer limit")
            parsed[name] = item
    return IngestLimits(**parsed)


def _source_from_request(source_fd: int, request: ParserRequest) -> ReadOnlySource:
    metadata = request.payload["source"]
    assert isinstance(metadata, Mapping)
    assert request.source_sha256 is not None
    return ReadOnlySource(
        fd=source_fd,
        source_sha256=request.source_sha256,
        source_id=metadata.get("source_id"),
        source_version=metadata.get("source_version"),
        filename=str(metadata["filename"]),
        mime_type=metadata.get("mime_type"),
    )


def _validate_source_metadata(value: object) -> None:
    if not isinstance(value, Mapping) or set(value) != _SOURCE_KEYS:
        raise SandboxProtocolError("parser source metadata does not match schema")
    filename = value.get("filename")
    if (
        not isinstance(filename, str)
        or not filename
        or len(filename) > 255
        or "/" in filename
        or "\\" in filename
        or filename in {".", ".."}
    ):
        raise SandboxProtocolError("parser filename must be a bounded basename")
    mime_type = value.get("mime_type")
    if mime_type is not None and (not isinstance(mime_type, str) or len(mime_type) > 255):
        raise SandboxProtocolError("invalid parser MIME type")
    source_id = value.get("source_id")
    if source_id is not None and (not isinstance(source_id, str) or len(source_id) > 128):
        raise SandboxProtocolError("invalid parser source identifier")
    source_version = value.get("source_version")
    if source_version is not None and (
        not isinstance(source_version, int)
        or isinstance(source_version, bool)
        or source_version < 1
    ):
        raise SandboxProtocolError("invalid parser source version")


def _validate_source_descriptor(source_fd: int, limits: IngestLimits) -> None:
    flags = fcntl.fcntl(source_fd, fcntl.F_GETFL)
    if flags & os.O_ACCMODE != os.O_RDONLY:
        raise SandboxProtocolError("parser source descriptor must be read-only")
    descriptor_stat = os.fstat(source_fd)
    if not stat.S_ISREG(descriptor_stat.st_mode):
        raise SandboxProtocolError("parser source descriptor must be a regular file")
    if descriptor_stat.st_size < 0 or descriptor_stat.st_size > limits.max_upload_bytes:
        raise SandboxProtocolError("parser source descriptor exceeds upload limit")


def _descriptors_from_ancillary(
    ancillary: list[tuple[int, int, bytes]],
) -> list[int]:
    descriptors: list[int] = []
    try:
        for level, kind, data in ancillary:
            if level != socket.SOL_SOCKET or kind != socket.SCM_RIGHTS:
                raise SandboxProtocolError("unsupported parser ancillary data")
            usable = len(data) - (len(data) % _FD_BYTES)
            values = array.array("i")
            values.frombytes(data[:usable])
            descriptors.extend(values.tolist())
        return descriptors
    except BaseException:
        for descriptor in descriptors:
            os.close(descriptor)
        raise


def _run_isolated_startup_self_test(
    entry: _AdapterEntry,
    limits: IngestLimits,
    config: ParserServiceConfig,
) -> StartupSelfTestResult:
    """Run startup evidence in a scrubbed, bounded, killable child."""

    outcome = _execute_child(
        lambda: _startup_self_test_payload(entry, limits),
        config,
        wall_timeout_seconds=min(config.wall_timeout_seconds, 5.0),
    )
    payload = outcome.payload
    if outcome.reason is not None or not isinstance(payload, Mapping):
        return StartupSelfTestResult(entry.format_key, False)
    if set(payload) != {"format_key", "passed", "evidence_digest"}:
        return StartupSelfTestResult(entry.format_key, False)
    if payload.get("format_key") != entry.format_key or payload.get("passed") is not True:
        return StartupSelfTestResult(entry.format_key, False)
    evidence_digest = payload.get("evidence_digest")
    if not _is_sha256(evidence_digest):
        return StartupSelfTestResult(entry.format_key, False)
    return StartupSelfTestResult(entry.format_key, True, str(evidence_digest))


def _startup_self_test_payload(
    entry: _AdapterEntry,
    limits: IngestLimits,
) -> Mapping[str, Any]:
    result = _run_startup_self_test(entry, limits)
    return {
        "format_key": result.format_key,
        "passed": result.passed,
        "evidence_digest": result.evidence_digest,
    }


def _run_startup_self_test(
    entry: _AdapterEntry,
    limits: IngestLimits,
) -> StartupSelfTestResult:
    """Return closed evidence from a real safe-fixture parse and preflight."""

    if entry.startup_fixture is None:
        return StartupSelfTestResult(entry.format_key, False)
    try:
        raw = entry.startup_fixture()
        if not isinstance(raw, bytes) or not raw or len(raw) > limits.max_upload_bytes:
            raise ValueError("startup fixture must be bounded non-empty bytes")
        source_sha256 = hashlib.sha256(raw).hexdigest()
        with tempfile.TemporaryFile() as handle:
            handle.write(raw)
            handle.flush()
            source = ReadOnlySource(
                handle.fileno(),
                source_sha256,
                filename=_startup_filename(entry.format_key),
                mime_type=MIME_BY_FILE_TYPE[FileType(entry.format_key)],
            )
            context = AdapterContext(
                adapter_key=entry.adapter_key,
                adapter_version=entry.adapter_version,
                policy_version="startup-self-test-v1",
                metadata={
                    "detected_type": entry.format_key,
                    "canonical_mime": MIME_BY_FILE_TYPE[FileType(entry.format_key)],
                },
            )
            result = entry.run(source, context, limits)
            if isinstance(result, AdapterRunResult):
                normalized = result.normalized
                generated = result.artifacts
            elif isinstance(result, NormalizedDocument):
                normalized = result
                generated = ()
            elif isinstance(result, Mapping):
                normalized = NormalizedDocument.model_validate(result)
                generated = ()
            else:
                raise TypeError("startup parser returned an unsupported result")
            if normalized.metadata.detected_type.value != entry.format_key:
                raise ValueError("startup parser returned the wrong format")
            normalized_payload = normalized.model_dump(mode="json")
            _require_bounded_json(
                "startup normalized output",
                normalized_payload,
                limits.max_normalized_output_bytes,
            )
            if sum(len(artifact.data) for artifact in generated) > (
                limits.max_normalized_output_bytes
            ):
                raise ResourceLimitExceeded(
                    "max_normalized_output_bytes",
                    limits.max_normalized_output_bytes,
                    sum(len(artifact.data) for artifact in generated),
                )
            parsed_artifacts = tuple(
                ParserArtifact(
                    descriptor_for_generated(
                        artifact,
                        ordinal=ordinal,
                        source_sha256=source_sha256,
                        seal_supported=False,
                        sealed=False,
                    ),
                    artifact.data,
                )
                for ordinal, artifact in enumerate(generated, start=1)
            )
            validate_result_artifact_set(
                entry.adapter_key,
                normalized,
                parsed_artifacts,
                source_sha256=source_sha256,
            )
            preflight = entry.preflight(
                source,
                {
                    "format": entry.format_key,
                    "canonical_mime": MIME_BY_FILE_TYPE[FileType(entry.format_key)],
                },
                limits,
            )
            if preflight.get("safe") is not True:
                raise ValueError("startup preflight did not return safe evidence")
        evidence = {
            "schema": "clerksan.parser-format-self-test",
            "version": 1,
            "adapter_key": entry.adapter_key,
            "adapter_version": entry.adapter_version,
            "format_key": entry.format_key,
            "fixture_sha256": source_sha256,
            "normalized_sha256": hashlib.sha256(_canonical_json(normalized_payload)).hexdigest(),
            "artifact_sha256": [
                hashlib.sha256(artifact.data).hexdigest() for artifact in generated
            ],
            "preflight_sha256": hashlib.sha256(_canonical_json(preflight)).hexdigest(),
        }
        return StartupSelfTestResult(
            entry.format_key,
            True,
            hashlib.sha256(_canonical_json(evidence)).hexdigest(),
        )
    except Exception:
        return StartupSelfTestResult(entry.format_key, False)


def _startup_fixture(format_key: str) -> bytes:
    """Build a tiny synthetic, non-personal safe fixture for one static format."""

    text_fixtures = {
        "md": b"# Safe startup fixture\n",
        "txt": b"safe startup fixture\n",
        "rst": b"Safe startup fixture\n====================\n",
        "log": b"INFO safe startup fixture\n",
        "csv": b"name,value\nsafe,1\n",
        "tsv": b"name\tvalue\nsafe\t1\n",
        "json": b'{"safe":true}',
        "jsonl": b'{"safe":true}\n',
        "yaml": b"safe: true\n",
        "xml": b"<root><value>safe</value></root>",
        "html": b"<!doctype html><html><body><p>safe</p></body></html>",
        "svg": b'<svg xmlns="http://www.w3.org/2000/svg"><title>safe</title></svg>',
        "rtf": b"{\\rtf1\\ansi Safe startup fixture}",
        "eml": (
            b"From: sender@example.invalid\r\n"
            b"To: receiver@example.invalid\r\n"
            b"Subject: Safe fixture\r\n"
            b"Content-Type: text/plain; charset=utf-8\r\n\r\n"
            b"safe startup fixture\r\n"
        ),
    }
    if format_key in text_fixtures:
        return text_fixtures[format_key]
    if format_key in {"jpeg", "png", "webp", "bmp", "gif", "tiff"}:
        from PIL import Image

        output = io.BytesIO()
        image = Image.new("RGB", (3, 2), "white")
        try:
            image.save(
                output,
                format={
                    "jpeg": "JPEG",
                    "png": "PNG",
                    "webp": "WEBP",
                    "bmp": "BMP",
                    "gif": "GIF",
                    "tiff": "TIFF",
                }[format_key],
            )
        finally:
            image.close()
        return output.getvalue()
    if format_key == "pdf":
        from pypdf import PdfWriter
        from pypdf.generic import (
            DecodedStreamObject,
            DictionaryObject,
            NameObject,
        )

        writer = PdfWriter()
        page = writer.add_blank_page(width=320, height=72)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
                NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
        )
        content = DecodedStreamObject()
        content.set_data(
            b"BT /F1 10 Tf 5 32 Td (safe bounded parser startup fixture text layer) Tj ET\n"
        )
        page.replace_contents(content)
        output = io.BytesIO()
        writer.write(output)
        return output.getvalue()
    if format_key == "docx":
        from docx import Document

        document = Document()
        document.add_paragraph("safe startup fixture")
        document.core_properties.created = _STARTUP_FIXTURE_DATETIME
        document.core_properties.modified = _STARTUP_FIXTURE_DATETIME
        output = io.BytesIO()
        document.save(output)
        return _deterministic_zip_bytes(output.getvalue())
    if format_key == "xlsx":
        from openpyxl import Workbook

        workbook = Workbook()
        try:
            workbook.active.append(["name", "value"])
            workbook.active.append(["safe", 1])
            workbook.properties.created = _STARTUP_FIXTURE_DATETIME
            workbook.properties.modified = _STARTUP_FIXTURE_DATETIME
            output = io.BytesIO()
            workbook.save(output)
            return _deterministic_zip_bytes(output.getvalue())
        finally:
            workbook.close()
    if format_key == "pptx":
        return _startup_pptx_fixture()
    if format_key in {"odt", "odp", "ods"}:
        return _startup_odf_fixture(format_key)
    if format_key == "zip":
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("safe.txt", "safe startup fixture")
        return _deterministic_zip_bytes(output.getvalue())
    if format_key in {"tar", "tgz"}:
        raw_tar = _startup_tar_fixture()
        return gzip.compress(raw_tar, mtime=0) if format_key == "tgz" else raw_tar
    if format_key == "gz":
        return gzip.compress(b"safe startup fixture", mtime=0)
    raise ValueError("no built-in startup fixture for parser format")


def _startup_filename(format_key: str) -> str:
    return {
        "jpeg": "fixture.jpg",
        "tiff": "fixture.tiff",
        "yaml": "fixture.yaml",
        "jsonl": "fixture.jsonl",
        "tgz": "fixture.tgz",
        "gz": "fixture.txt.gz",
    }.get(format_key, f"fixture.{format_key}")


def _startup_tar_fixture() -> bytes:
    output = io.BytesIO()
    data = b"safe startup fixture"
    with tarfile.open(fileobj=output, mode="w") as archive:
        member = tarfile.TarInfo("safe.txt")
        member.size = len(data)
        member.mtime = 0
        archive.addfile(member, io.BytesIO(data))
    return output.getvalue()


def _deterministic_zip_bytes(raw: bytes) -> bytes:
    """Remove timestamps and mutable ZIP metadata from synthetic probe fixtures."""

    output = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(raw), "r") as source,
        zipfile.ZipFile(
            output,
            "w",
        ) as target,
    ):
        for source_info in source.infolist():
            member = source.read(source_info)
            if source_info.filename == "docProps/core.xml":
                for field in (b"created", b"modified"):
                    pattern = (
                        rb"(<dcterms:" + field + rb"\b[^>]*>)[^<]*(</dcterms:" + field + rb">)"
                    )
                    member = re.sub(
                        pattern,
                        rb"\g<1>2000-01-01T00:00:00Z\g<2>",
                        member,
                    )
            info = zipfile.ZipInfo(
                source_info.filename,
                date_time=_STARTUP_FIXTURE_ZIP_TIMESTAMP,
            )
            info.compress_type = source_info.compress_type
            info.external_attr = source_info.external_attr
            info.create_system = source_info.create_system
            target.writestr(info, member)
    return output.getvalue()


def _startup_pptx_fixture() -> bytes:
    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_relationship_namespace = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    )
    content_types = (
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.'
        'openxmlformats-officedocument.presentationml.presentation.main+xml"/>'
        "</Types>"
    )
    presentation = (
        f'<p:presentation xmlns:p="urn:p" xmlns:r="{office_relationship_namespace}">'
        '<p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>'
        "</p:presentation>"
    )
    relationships = (
        f'<Relationships xmlns="{relationship_namespace}"><Relationship Id="rId1" '
        f'Type="{office_relationship_namespace}/slide" Target="slides/slide1.xml"/>'
        "</Relationships>"
    )
    slide = (
        '<p:sld xmlns:p="urn:p" xmlns:a="urn:a"><p:cSld><p:spTree><p:sp>'
        "<p:txBody><a:p><a:r><a:t>safe startup fixture</a:t></a:r></a:p>"
        "</p:txBody></p:sp></p:spTree></p:cSld></p:sld>"
    )
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("ppt/presentation.xml", presentation)
        archive.writestr("ppt/_rels/presentation.xml.rels", relationships)
        archive.writestr("ppt/slides/slide1.xml", slide)
    return _deterministic_zip_bytes(output.getvalue())


def _startup_odf_fixture(format_key: str) -> bytes:
    mime = {
        "odt": "application/vnd.oasis.opendocument.text",
        "odp": "application/vnd.oasis.opendocument.presentation",
        "ods": "application/vnd.oasis.opendocument.spreadsheet",
    }[format_key]
    body = {
        "odt": "<office:body><office:text><text:p>safe</text:p></office:text></office:body>",
        "odp": (
            '<office:body><office:presentation><draw:page draw:name="safe">'
            "<text:p>safe</text:p></draw:page></office:presentation></office:body>"
        ),
        "ods": (
            '<office:body><office:spreadsheet><table:table table:name="safe">'
            "<table:table-row><table:table-cell><text:p>value</text:p>"
            "</table:table-cell></table:table-row></table:table>"
            "</office:spreadsheet></office:body>"
        ),
    }[format_key]
    namespaces = (
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0" '
        'xmlns:xlink="http://www.w3.org/1999/xlink"'
    )
    content = f"<office:document-content {namespaces}>{body}</office:document-content>"
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("mimetype", mime, compress_type=zipfile.ZIP_STORED)
        archive.writestr("content.xml", content)
        archive.writestr("META-INF/manifest.xml", '<manifest xmlns="urn:manifest"/>')
    return _deterministic_zip_bytes(output.getvalue())


def _delimited_handler(format_key: str) -> RunHandler:
    def normalize(
        source: ReadOnlySource, context: AdapterContext, limits: IngestLimits
    ) -> NormalizedDocument:
        from clerksan.ingest.adapters.delimited import DelimitedAdapter

        metadata = dict(context.metadata)
        detected_type = metadata.get("detected_type")
        if detected_type not in {None, format_key}:
            raise ValueError("adapter metadata conflicts with registered format")
        metadata["detected_type"] = format_key
        bound_context = AdapterContext(
            adapter_key=context.adapter_key,
            adapter_version=context.adapter_version,
            policy_version=context.policy_version,
            registry_digest=context.registry_digest,
            metadata=metadata,
        )
        return DelimitedAdapter(limits=limits).normalize(source, bound_context)

    return normalize


@dataclass(frozen=True, slots=True)
class _ParserPdfPolicy:
    """Secret-free PDF thresholds owned by the parser image."""

    pdf_min_chars_per_page: int = 32
    pdf_mojibake_ratio: float = 0.30


class _NoOcr:
    """Make accidental model use inside the structural sidecar fail closed."""

    async def extract_markdown(self, *_args: Any, **_kwargs: Any) -> str:
        raise RuntimeError("OCR is unavailable inside the structural parser sidecar")


def _adapter_handler(
    format_key: str,
    factory: Callable[[IngestLimits], Any],
) -> RunHandler:
    """Bind one static adapter factory to exactly one advertised format key."""

    def normalize(
        source: ReadOnlySource,
        context: AdapterContext,
        limits: IngestLimits,
    ) -> NormalizedDocument | AdapterRunResult:
        metadata = dict(context.metadata)
        detected_type = metadata.get("detected_type")
        if detected_type not in {None, format_key}:
            raise ValueError("adapter metadata conflicts with registered format")
        metadata["detected_type"] = format_key
        bound_context = AdapterContext(
            adapter_key=context.adapter_key,
            adapter_version=context.adapter_version,
            policy_version=context.policy_version,
            registry_digest=context.registry_digest,
            metadata=metadata,
        )
        adapter = factory(limits)
        if hasattr(adapter, "normalize_with_artifacts"):
            result = adapter.normalize_with_artifacts(source, bound_context)
        else:
            result = adapter.normalize(source, bound_context)
        normalized = result.normalized if isinstance(result, AdapterRunResult) else result
        if normalized.metadata.detected_type.value != format_key:
            raise ValueError("adapter normalized a different format than it advertised")
        return result

    return normalize


def _normalizing_preflight(format_key: str, handler: RunHandler) -> PreflightHandler:
    """Run the same bounded parser before persistence without returning private content."""

    def preflight(
        source: ReadOnlySource,
        detected: Mapping[str, Any],
        limits: IngestLimits,
    ) -> Mapping[str, Any]:
        if detected.get("format") != format_key:
            raise ValueError("preflight format conflicts with registered adapter")
        result = handler(
            source,
            AdapterContext(
                adapter_key=(
                    f"delimited.{format_key}" if format_key in {"csv", "tsv"} else format_key
                ),
                policy_version="universal-intake-v1",
                metadata={
                    "detected_type": format_key,
                    "canonical_mime": str(detected.get("canonical_mime") or ""),
                },
            ),
            limits,
        )
        normalized = result.normalized if isinstance(result, AdapterRunResult) else result
        if not isinstance(normalized, NormalizedDocument):
            raise TypeError("preflight parser returned an unsupported result")
        normalized_payload = normalized.model_dump(mode="json")
        _require_bounded_json(
            "normalized output",
            normalized_payload,
            limits.max_normalized_output_bytes,
        )
        return {
            "safe": True,
            "detected_format": format_key,
            "policy": "bounded-normalizing-preflight-v1",
            "normalized_sha256": hashlib.sha256(_canonical_json(normalized_payload)).hexdigest(),
            "table_count": len(normalized.tables),
            "image_count": len(normalized.images),
        }

    return preflight


def _safe_regular_file_preflight(
    source: ReadOnlySource, detected: Mapping[str, Any], limits: IngestLimits
) -> Mapping[str, Any]:
    _validate_source_descriptor(source.fd, limits)
    source.verify_digest(max_bytes=limits.max_upload_bytes)
    return {
        "safe": True,
        "detected_format": str(detected["format"]),
        "policy": "bounded-regular-file-v1",
    }


def _apply_child_limits(config: ParserServiceConfig) -> None:
    _set_resource_limit("RLIMIT_CORE", 0)
    _set_resource_limit("RLIMIT_CPU", config.cpu_seconds)
    _set_resource_limit("RLIMIT_AS", config.memory_bytes)
    _set_resource_limit("RLIMIT_DATA", config.memory_bytes)
    _set_resource_limit("RLIMIT_FSIZE", config.file_bytes)
    _set_resource_limit("RLIMIT_NOFILE", config.open_files)
    _set_resource_limit("RLIMIT_NPROC", config.processes)


def _set_resource_limit(name: str, requested: int) -> None:
    limit_kind = getattr(resource, name, None)
    if limit_kind is None:
        return
    _soft, hard = resource.getrlimit(limit_kind)
    bounded = requested if hard == resource.RLIM_INFINITY else min(requested, hard)
    try:
        resource.setrlimit(limit_kind, (bounded, bounded))
    except ValueError:
        # macOS refuses to lower some hard VM limits below the process's current
        # mapped address range. Keep the inherited hard ceiling there and apply
        # the requested soft bound; Linux release containers take the hard path.
        try:
            resource.setrlimit(limit_kind, (bounded, hard))
        except ValueError:
            # The release sandbox is additionally hard-bounded by its cgroup and
            # refuses verification without that evidence. This compatibility
            # branch lets the host-only test harness exercise kill/reap behavior.
            return


def _close_unneeded_descriptors(keep: set[int]) -> None:
    try:
        descriptors = [int(name) for name in os.listdir("/proc/self/fd")]
    except (FileNotFoundError, ValueError):
        upper = resource.getrlimit(resource.RLIMIT_NOFILE)[0]
        descriptors = list(range(3, min(int(upper), 1024)))
    for descriptor in descriptors:
        if descriptor not in keep:
            try:
                os.close(descriptor)
            except OSError:
                pass


def _kill_and_reap(pid: int) -> bool:
    try:
        if os.getpgid(pid) == pid:
            os.killpg(pid, signal.SIGKILL)
        else:
            os.kill(pid, signal.SIGKILL)
    except ProcessLookupError:
        pass
    except OSError:
        try:
            os.kill(pid, signal.SIGKILL)
        except ProcessLookupError:
            pass
    while True:
        try:
            waited_pid, _ = os.waitpid(pid, 0)
            return waited_pid == pid
        except InterruptedError:
            continue
        except ChildProcessError:
            return True


def _write_all(descriptor: int, data: bytes) -> None:
    offset = 0
    while offset < len(data):
        offset += os.write(descriptor, data[offset:])


def _send_response(
    connection: socket.socket,
    payload: Mapping[str, Any],
    descriptors: tuple[int, ...] | list[int] = (),
) -> None:
    encoded = _canonical_json(payload) + b"\n"
    if len(encoded) > MAX_PROTOCOL_RESPONSE_BYTES:
        raise SandboxProtocolError("parser response exceeds protocol limit")
    _send_encoded_response(connection, encoded, descriptors)


def _send_encoded_response(
    connection: socket.socket,
    encoded: bytes,
    descriptors: tuple[int, ...] | list[int],
) -> None:
    if len(descriptors) > MAX_ARTIFACT_FDS:
        raise SandboxProtocolError("too many parser artifact descriptors")
    if not descriptors:
        connection.sendall(encoded)
        return
    rights = array.array("i", descriptors)
    sent = connection.sendmsg(
        [encoded],
        [(socket.SOL_SOCKET, socket.SCM_RIGHTS, rights.tobytes())],
    )
    if sent < 1:
        raise SandboxProtocolError("parser artifact response send made no progress")
    if sent < len(encoded):
        connection.sendall(encoded[sent:])


def _close_descriptors(descriptors) -> None:
    for descriptor in descriptors:
        try:
            os.close(descriptor)
        except OSError:
            pass


def _canonical_json(value: object) -> bytes:
    try:
        return json.dumps(value, sort_keys=True, separators=(",", ":"), ensure_ascii=False).encode(
            "utf-8"
        )
    except (TypeError, ValueError, UnicodeEncodeError) as error:
        raise SandboxProtocolError("parser protocol value is not JSON-safe") from error


def _require_bounded_json(name: str, value: object, maximum: int) -> None:
    if len(_canonical_json(value)) > maximum:
        raise SandboxProtocolError(f"{name} exceeds protocol limit")


def _require_stable_token(name: str, value: object) -> None:
    if (
        not isinstance(value, str)
        or not value
        or len(value) > 128
        or any(not (char.isascii() and (char.isalnum() or char in "._+-")) for char in value)
    ):
        raise SandboxProtocolError(f"invalid {name}")


def _reject_forbidden_request_data(value: object) -> None:
    if isinstance(value, Mapping):
        for key, child in value.items():
            if not isinstance(key, str):
                raise SandboxProtocolError("parser request keys must be strings")
            normalized = key.casefold().replace("-", "_")
            if normalized in _FORBIDDEN_REQUEST_KEYS or normalized.endswith(
                ("_password", "_secret", "_token", "_path")
            ):
                raise SandboxProtocolError(
                    "host paths and secrets are forbidden in parser requests"
                )
            _reject_forbidden_request_data(child)
    elif isinstance(value, (list, tuple)):
        for child in value:
            _reject_forbidden_request_data(child)
    elif isinstance(value, str) and (
        value.startswith(("/", "\\"))
        or "://" in value
        or (len(value) >= 3 and value[1] == ":" and value[2] in "/\\")
    ):
        raise SandboxProtocolError("host paths and remote URLs are forbidden in parser requests")


def _is_sha256(value: object) -> bool:
    return (
        isinstance(value, str)
        and len(value) == 64
        and all(char in "0123456789abcdef" for char in value)
    )


def _sleep_then_payload(seconds: float) -> Mapping[str, Any]:
    time.sleep(seconds)
    return {"completed": True}


def _network_namespace_isolated() -> bool:
    try:
        routes = Path("/proc/net/route").read_text(encoding="ascii").splitlines()[1:]
    except (OSError, UnicodeError):
        return False
    return not any(
        fields[1] == "00000000" and fields[7] == "00000000"
        for line in routes
        if len(fields := line.split()) >= 8
    )


def _application_secrets_absent() -> bool:
    for name in os.environ:
        upper = name.upper()
        if any(fragment in upper for fragment in _FORBIDDEN_ENVIRONMENT_FRAGMENTS):
            return False
    return True


def _status_value(name: str) -> str | None:
    try:
        lines = Path("/proc/self/status").read_text(encoding="ascii").splitlines()
    except (OSError, UnicodeError):
        return None
    prefix = f"{name}:"
    for line in lines:
        if line.startswith(prefix):
            return line.split(":", 1)[1].strip()
    return None


def _mount_entry(path: str) -> tuple[str, frozenset[str]] | None:
    try:
        lines = Path("/proc/self/mountinfo").read_text(encoding="utf-8").splitlines()
    except (OSError, UnicodeError):
        return None
    for line in lines:
        fields = line.split()
        if len(fields) < 10 or fields[4] != path or "-" not in fields:
            continue
        separator = fields.index("-")
        filesystem = fields[separator + 1]
        options = frozenset(fields[5].split(",")) | frozenset(fields[separator + 3].split(","))
        return filesystem, options
    return None


def _mount_is_read_only(path: str) -> bool:
    entry = _mount_entry(path)
    return entry is not None and "ro" in entry[1]


def _tmpfs_is_hardened(path: str) -> bool:
    entry = _mount_entry(path)
    if entry is None:
        return False
    filesystem, options = entry
    return filesystem == "tmpfs" and {"nodev", "noexec", "nosuid"}.issubset(options)


def _cgroup_is_bounded(config: ParserServiceConfig) -> bool:
    try:
        memory = int(Path("/sys/fs/cgroup/memory.max").read_text(encoding="ascii").strip())
        processes = int(Path("/sys/fs/cgroup/pids.max").read_text(encoding="ascii").strip())
        quota_raw, period_raw = (
            Path("/sys/fs/cgroup/cpu.max").read_text(encoding="ascii").strip().split()
        )
        quota = int(quota_raw)
        period = int(period_raw)
    except (OSError, UnicodeError, ValueError):
        return False
    return (
        memory <= config.memory_bytes + 128 * 1024 * 1024
        and processes <= max(16, config.processes * 4)
        and quota > 0
        and period > 0
        and quota / period <= 1.0
    )


def _socket_is_permissioned(path: Path, expected_mode: int) -> bool:
    try:
        socket_stat = os.lstat(path)
    except OSError:
        return False
    return (
        stat.S_ISSOCK(socket_stat.st_mode)
        and socket_stat.st_uid == os.geteuid()
        and stat.S_IMODE(socket_stat.st_mode) == expected_mode
        and socket_stat.st_mode & 0o007 == 0
    )


def _unlink_owned_socket(path: Path) -> None:
    try:
        existing = os.lstat(path)
    except FileNotFoundError:
        return
    if not stat.S_ISSOCK(existing.st_mode) or existing.st_uid != os.geteuid():
        raise RuntimeError("refusing to replace an untrusted parser socket path")
    path.unlink()


def _main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Clerk-san isolated parser sidecar")
    subcommands = parser.add_subparsers(dest="command", required=True)
    subcommands.add_parser("serve", help="serve parser requests")
    subcommands.add_parser("probe", help="probe a running sidecar")
    args = parser.parse_args(argv)
    config = ParserServiceConfig.from_environment()
    if args.command == "serve":
        ParserSidecarServer(config).serve_forever()
        return 0
    result = SidecarSandboxBackend(
        str(config.socket_path), timeout_seconds=config.wall_timeout_seconds + 2
    ).startup_probe()
    return 0 if result.verified else 1


if __name__ == "__main__":
    sys.exit(_main())
