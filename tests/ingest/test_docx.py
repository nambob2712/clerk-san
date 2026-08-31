from __future__ import annotations

import hashlib
import io
import zipfile

import pytest
from docx import Document
from PIL import Image

from clerksan.ingest.adapters.docx import DocxAdapter, extract_embedded_images
from clerksan.ingest.filetype import FileType
from clerksan.ingest.limits import UnsafeArchiveMemberError
from clerksan.ingest.normalized import DocMetadata


def _png(width: int, height: int, *, color: str) -> bytes:
    image = Image.new("RGB", (width, height), color=color)
    output = io.BytesIO()
    image.save(output, format="PNG")
    image.close()
    return output.getvalue()


def _document_bytes() -> bytes:
    document = Document()
    document.add_heading("経費報告", level=1)
    document.add_paragraph("7月分の精算です。")
    table = document.add_table(rows=3, cols=4)
    values = [
        ["日付", "取引先", "金額", "備考"],
        ["2026-07-01", "サンプル商店", "1200", "交通費"],
        ["2026-07-02", "書店", "800", "資料"],
    ]
    for row, values_row in zip(table.rows, values, strict=True):
        for cell, value in zip(row.cells, values_row, strict=True):
            cell.text = value
    document.add_picture(io.BytesIO(_png(60, 60, color="navy")))
    document.add_picture(io.BytesIO(_png(40, 40, color="gray")))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _metadata(raw: bytes) -> DocMetadata:
    return DocMetadata(
        filename="expense-report.docx",
        detected_type=FileType.DOCX,
        sha256=hashlib.sha256(raw).hexdigest(),
        extra={"content_path": "documents/expense-report.docx"},
    )


async def test_docx_adapter_preserves_headings_tables_and_embedded_image_linkage() -> None:
    raw = _document_bytes()
    result = await DocxAdapter().adapt(raw, _metadata(raw))

    assert "# 経費報告" in result.markdown_body
    assert "| 日付 | 取引先 | 金額 | 備考 |" in result.markdown_body
    assert len(result.tables) == 1
    assert result.tables[0].header == ["日付", "取引先", "金額", "備考"]
    assert result.tables[0].rows == [
        ["2026-07-01", "サンプル商店", "1200", "交通費"],
        ["2026-07-02", "書店", "800", "資料"],
    ]
    assert len(result.images) == 1
    assert result.images[0].source_location == "word/media/image1.png"
    assert result.images[0].content_path.startswith("embedded/sha256/")
    assert result.metadata.extra["embedded_image_count"] == 1


def test_docx_embedded_images_skip_small_media_and_are_content_addressed() -> None:
    raw = _document_bytes()
    images = extract_embedded_images(raw)

    assert len(images) == 1
    assert images[0].width == 60
    assert images[0].height == 60
    assert images[0].sha256 in images[0].content_path


def test_docx_embedded_images_deduplicate_identical_media_by_checksum() -> None:
    output = io.BytesIO()
    image = _png(60, 60, color="purple")
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("word/media/image1.png", image)
        archive.writestr("word/media/image2.png", image)

    images = extract_embedded_images(output.getvalue())

    assert len(images) == 1
    assert images[0].source_location == "word/media/image1.png"


async def test_docx_rejects_unsafe_archive_member_before_parser_opens_it() -> None:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("../outside.xml", "not a document")
    raw = output.getvalue()

    with pytest.raises(UnsafeArchiveMemberError):
        await DocxAdapter().adapt(raw, _metadata(raw))
